"""
Tạo 4 file .docx hướng dẫn từng Task cho dự án Smart Factory Camera Module.
Chạy: python generate_task_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "HUONGDAN")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def add_step(doc, number, title, details):
    p = doc.add_paragraph()
    run_num = p.add_run(f"Bước {number}: ")
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)
    for detail in details:
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(detail).font.size = Pt(11)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠ Lưu ý: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x55, 0x00)
    p.add_run(text).font.size = Pt(11)


def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.left_indent = Inches(0.5)


def header_page(doc, task_num, task_title, subtitle):
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"TASK {task_num}")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run(task_title)
    run2.bold = True
    run2.font.size = Pt(18)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = t3.add_run(subtitle)
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph("─" * 70)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────
# TASK 1: BP_StateManager
# ─────────────────────────────────────────────────────────
def make_task1():
    doc = Document()
    header_page(doc, 1, "BP_StateManager",
                "Blueprint quản lý trạng thái | HTTP poll localhost:8080 | 6 states | Event Dispatcher | Phím 1–6")

    set_heading(doc, "Tổng quan", 1)
    doc.add_paragraph(
        "BP_StateManager là Blueprint Actor trung tâm của toàn bộ module. "
        "Nó poll dữ liệu từ Python server mỗi 3 giây, phân loại trạng thái, "
        "broadcast event OnStateChanged để các BP khác (đèn, HUD) phản ứng, "
        "và cho phép tester ép buộc state bằng phím 1–6."
    )

    set_heading(doc, "Yêu cầu trước khi làm", 2)
    doc.add_paragraph("Hoàn thành các mục sau trước khi bắt đầu:")
    for item in [
        "UE5 project SmartFactory_CameraModule mở được (đã compile C++ thành công).",
        "Python server đang chạy: python camera_server.py  →  http://localhost:8080",
        "Plugin VaRest đã bật: Edit → Plugins → tìm VaRest → Enable → Restart Editor.",
        "Đã có enum EModuleState trong C++ (ScanningDataTypes.h) và UE đã nhận diện enum này.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_note(doc, "Nếu không thấy VaRest trong Plugins, cài từ Fab/Marketplace hoặc "
                  "dùng cách thay thế: tạo Blueprint con của AStateManagerActor (C++ đã viết sẵn HTTP).")

    # BƯỚC 1
    set_heading(doc, "PHẦN A — Tạo Blueprint và thêm biến", 1)
    add_step(doc, 1, "Tạo asset BP_StateManager", [
        "Mở Content Browser → Content/CameraModule/Core.",
        "Right-click vùng trống → Blueprint Class.",
        "Parent class: chọn Actor → bấm Select.",
        "Đặt tên: BP_StateManager → Enter.",
        "Double-click để mở Blueprint Editor.",
    ])

    add_step(doc, 2, "Thêm biến (My Blueprint → VARIABLES → bấm +)", [
        "CurrentState    |  kiểu: EModuleState  |  default: Idle",
        "ServerUrl       |  kiểu: String        |  default: http://127.0.0.1:8080/current",
        "PollInterval    |  kiểu: Float         |  default: 3.0",
        "CountOK         |  kiểu: Integer       |  default: 0",
        "CountWarning    |  kiểu: Integer       |  default: 0",
        "CountError      |  kiểu: Integer       |  default: 0",
        "LastObjectId    |  kiểu: String        |  default: (trống)",
        "LastStatus      |  kiểu: String        |  default: (trống)",
        "LastTemperature |  kiểu: Float         |  default: 0.0",
    ])
    add_note(doc, "Sau khi đặt kiểu EModuleState, bấm Compile để thấy Default Value dropdown xuất hiện.")

    # BƯỚC 3 - Event Dispatcher
    set_heading(doc, "PHẦN B — Tạo Event Dispatcher OnStateChanged", 1)
    add_step(doc, 3, "Tạo Event Dispatcher", [
        "My Blueprint → EVENT DISPATCHERS → bấm +.",
        "Đặt tên: OnStateChanged.",
        "Click chọn dispatcher vừa tạo → Details bên phải.",
        "Bấm + trong mục Inputs để thêm tham số:",
        "    Tên: NewState    |    Kiểu: EModuleState",
        "Compile + Save.",
    ])

    # BƯỚC 4 - SetModuleState function
    set_heading(doc, "PHẦN C — Hàm helper SetModuleState", 1)
    add_step(doc, 4, "Tạo function SetModuleState", [
        "My Blueprint → FUNCTIONS → bấm +.",
        "Đặt tên: SetModuleState.",
        "Details của function → INPUTS → bấm +: tên NewState, kiểu EModuleState.",
        "Mở graph function.",
        "Kéo CurrentState từ Variables vào graph → chọn Set.",
        "Nối pin NewState (từ Function Entry) vào input của Set CurrentState.",
        "Kéo OnStateChanged từ Event Dispatchers vào graph → chọn Call (Broadcast).",
        "Nối pin NewState vào tham số của node Call OnStateChanged.",
        "Nối exec: Function Entry → Set CurrentState → Call OnStateChanged.",
        "Compile + Save.",
    ])

    # BƯỚC 5 - BeginPlay + Timer
    set_heading(doc, "PHẦN D — BeginPlay + Timer tự động poll", 1)
    add_step(doc, 5, "Mở Event Graph và cấu hình BeginPlay", [
        "Click tab EventGraph.",
        "Tìm node Event BeginPlay (node đỏ sẵn có).",
        "Kéo từ exec pin của BeginPlay → tạo node Set Timer by Function Name:",
        "    Function Name: PollServer",
        "    Time: kéo từ biến PollInterval (Get)",
        "    Looping: tick true",
        "Compile + Save.",
    ])

    add_step(doc, 6, "Tạo function PollServer (gọi HTTP GET)", [
        "My Blueprint → FUNCTIONS → + → đặt tên: PollServer.",
        "Mở graph function PollServer.",
        "─── Nếu có VaRest ───",
        "Right-click → tìm 'VaRest' → chọn node Construct VaRest Request hoặc tương tự.",
        "Set URL: kéo biến ServerUrl (Get).",
        "Set Verb: GET.",
        "Gắn event On Request Complete để nhận response.",
        "─── Trong callback thành công ───",
        "Lấy field status (string): Get String Field, key='status'.",
        "Lấy field temperature_c (float): Get Number Field, key='temperature_c'.",
        "Lấy object_id, object_type → lưu vào LastObjectId, LastStatus.",
        "Dùng Branch + Switch để map status → EModuleState (xem Bước 7).",
        "─── Trong callback thất bại ───",
        "Gọi SetModuleState(Error).",
        "Print String: 'HTTP poll failed'.",
    ])

    add_step(doc, 7, "Map status string → EModuleState", [
        "Từ Get String Field (status), kéo output string ra.",
        "Right-click → tạo node Switch on String.",
        "Thêm các case:",
        "    'ok'      → SetModuleState(Detected_OK)    → CountOK + 1",
        "    'warning' → SetModuleState(Detected_Warning)→ CountWarning + 1",
        "    'error'   → SetModuleState(Error)           → CountError + 1",
        "    Default   → SetModuleState(Detected_Warning)",
        "Thêm check nhiệt độ trước switch: nếu temperature_c > 55 → SetModuleState(Fire_Alert).",
        "Compile + Save.",
    ])

    # BƯỚC 8 - Phím 1-6
    set_heading(doc, "PHẦN E — Phím 1–6 force state (debug)", 1)
    add_step(doc, 8, "Thêm keyboard debug input trong EventGraph", [
        "Mở tab EventGraph.",
        "Right-click → Keyboard Events → 1  (tạo event Key 1).",
        "Từ exec pin của Key 1 → gọi SetModuleState(Idle).",
        "Lặp lại tương tự:",
        "    Key 2 → SetModuleState(Scanning)",
        "    Key 3 → SetModuleState(Detected_OK)",
        "    Key 4 → SetModuleState(Detected_Warning)",
        "    Key 5 → SetModuleState(Error)",
        "    Key 6 → SetModuleState(Fire_Alert)",
        "(Tùy chọn) Thêm Print String sau mỗi SetModuleState để debug.",
        "Compile + Save.",
    ])

    # BƯỚC 9 - Đặt vào level
    set_heading(doc, "PHẦN F — Đặt vào level và test", 1)
    add_step(doc, 9, "Đặt BP_StateManager vào level", [
        "Mở Level Editor (tab map, không phải Blueprint Editor).",
        "Content Browser → Content/CameraModule/Core → kéo BP_StateManager vào Viewport.",
        "Chọn actor → Details → tìm Auto Receive Input → chọn Player 0.",
        "Ctrl+S lưu level.",
    ])

    add_step(doc, 10, "Chạy test", [
        "Mở terminal: python camera_server.py  (giữ cửa sổ mở).",
        "Bấm Play trong UE.",
        "Chờ 3 giây → phải thấy Print String hoặc state thay đổi theo mock data.",
        "Bấm phím 1 → state Idle.",
        "Bấm phím 6 → state Fire_Alert.",
        "Mở Output Log nếu cần kiểm tra lỗi HTTP.",
    ])

    set_heading(doc, "Checklist hoàn thành TASK 1", 1)
    for item in [
        "BP_StateManager tồn tại trong Content Browser.",
        "Có đầy đủ 6 biến (CurrentState, ServerUrl, PollInterval, Count*).",
        "Event Dispatcher OnStateChanged có tham số NewState: EModuleState.",
        "Function SetModuleState: set biến + broadcast dispatcher.",
        "Function PollServer: gọi HTTP → map status → gọi SetModuleState.",
        "BeginPlay gắn timer PollServer mỗi 3 giây, looping true.",
        "Phím 1–6 force state thành công.",
        "Actor trong level có Auto Receive Input = Player 0.",
        "Compile không lỗi đỏ.",
    ]:
        doc.add_paragraph(f"☐  {item}", style="List Bullet")

    doc.save(os.path.join(OUTPUT_DIR, "TASK1_BP_StateManager.docx"))
    print("✓ TASK1_BP_StateManager.docx")


# ─────────────────────────────────────────────────────────
# TASK 2: BP_CameraRig
# ─────────────────────────────────────────────────────────
def make_task2():
    doc = Document()
    header_page(doc, 2, "BP_CameraRig",
                "SpotLight đổi màu theo state | Scan beam hiện/ẩn khi Scanning | Bind OnStateChanged")

    set_heading(doc, "Tổng quan", 1)
    doc.add_paragraph(
        "BP_CameraRig là Blueprint Actor đại diện cho camera vật lý trong nhà máy. "
        "Nó lắng nghe event OnStateChanged từ BP_StateManager và phản hồi bằng cách "
        "đổi màu đèn SpotLight + hiện/ẩn beam quét theo từng state."
    )

    set_heading(doc, "Yêu cầu trước khi làm", 2)
    for item in [
        "TASK 1 hoàn thành: BP_StateManager đã có OnStateChanged dispatcher.",
        "BP_StateManager đã được đặt trong level.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    set_heading(doc, "PHẦN A — Tạo Blueprint và thêm Components", 1)
    add_step(doc, 1, "Tạo BP_CameraRig", [
        "Content Browser → Content/CameraModule/Actors.",
        "Right-click → Blueprint Class → Actor → đặt tên BP_CameraRig.",
        "Double-click mở Blueprint Editor.",
    ])

    add_step(doc, 2, "Thêm Components (panel Components → bấm + Add)", [
        "StaticMeshComponent  → đặt tên CameraBody  (dùng mesh camera hoặc cube tạm).",
        "SpotLightComponent   → đặt tên ScanLight   (đặt con của CameraBody).",
        "StaticMeshComponent  → đặt tên ScanBeam    (hình trụ/cone mỏng, đặt trước camera).",
    ])

    add_step(doc, 3, "Cấu hình ScanBeam ban đầu", [
        "Chọn ScanBeam trong Components.",
        "Details → Rendering → Hidden in Game: bật (hidden mặc định).",
        "Scale: (0.1, 0.1, 2.0) hoặc tùy ý để thể hiện tia quét.",
    ])

    set_heading(doc, "PHẦN B — Map màu theo EModuleState", 1)
    add_step(doc, 4, "Tạo function GetColorForState", [
        "My Blueprint → FUNCTIONS → + → đặt tên: GetColorForState.",
        "Thêm Input: InState kiểu EModuleState.",
        "Thêm Output: OutColor kiểu LinearColor.",
        "Mở graph → tạo Switch on EModuleState.",
        "Nối từng case với node Make LinearColor tương ứng:",
        "    Idle           → xanh dương nhạt  (R=0.0, G=0.5, B=1.0, A=1.0)",
        "    Scanning       → vàng              (R=1.0, G=0.9, B=0.0, A=1.0)",
        "    Detected_OK    → xanh lá           (R=0.0, G=1.0, B=0.1, A=1.0)",
        "    Detected_Warning→ cam              (R=1.0, G=0.4, B=0.0, A=1.0)",
        "    Error          → đỏ               (R=1.0, G=0.0, B=0.0, A=1.0)",
        "    Fire_Alert     → đỏ cam chớp      (R=1.0, G=0.1, B=0.0, A=1.0)",
        "Nối output của từng Make LinearColor → pin OutColor của Return Node.",
        "Compile + Save.",
    ])

    add_step(doc, 5, "Tạo function ApplyState (cập nhật đèn + beam)", [
        "My Blueprint → FUNCTIONS → + → đặt tên: ApplyState.",
        "Thêm Input: NewState kiểu EModuleState.",
        "Mở graph.",
        "─── Đổi màu đèn ───",
        "Gọi GetColorForState(NewState) → nhận OutColor.",
        "Kéo ScanLight vào graph → Set Light Color → truyền OutColor.",
        "─── Hiện/ẩn ScanBeam ───",
        "Tạo node Equal (EModuleState): NewState == Scanning.",
        "Kết quả bool → Set Hidden in Game cho ScanBeam:",
        "    true (đang Scanning) → Hidden = false (hiện beam)",
        "    false                → Hidden = true  (ẩn beam)",
        "Nối exec toàn bộ flow.",
        "Compile + Save.",
    ])

    set_heading(doc, "PHẦN C — Nhận event từ BP_StateManager", 1)
    add_step(doc, 6, "Bind OnStateChanged trong BeginPlay", [
        "Mở EventGraph.",
        "Từ Event BeginPlay, kéo exec ra.",
        "Tạo node Get All Actors Of Class, Class = BP_StateManager.",
        "Từ output array, tạo Get (index 0) → lấy actor BP_StateManager.",
        "Từ actor đó, kéo pin → tìm Bind Event to OnStateChanged.",
        "Tạo Custom Event tên HandleStateChanged, input: NewState EModuleState.",
        "Nối Custom Event vào Event pin của Bind node.",
        "Trong graph của HandleStateChanged:",
        "    Gọi ApplyState(NewState).",
        "Compile + Save.",
    ])

    add_note(doc, "Get All Actors Of Class chỉ hoạt động lúc runtime (trong Play). "
                  "Đảm bảo BP_StateManager đã có mặt trong level trước khi Play.")

    set_heading(doc, "PHẦN D — (Tùy chọn) Hiệu ứng Fire_Alert", 1)
    add_step(doc, 7, "Thêm chớp đèn khi Fire_Alert", [
        "Trong ApplyState, sau khi Set Light Color:",
        "Tạo Branch: NewState == Fire_Alert.",
        "Nhánh True: Set Timer by Function Name → FlashLight, Time=0.3, Looping=true.",
        "Nhánh False: Clear Timer by Function Name → FlashLight (dừng chớp).",
        "Tạo function FlashLight:",
        "    Lấy Light Intensity hiện tại.",
        "    Toggle giữa 5000 và 0 mỗi lần gọi.",
        "Compile + Save.",
    ])

    set_heading(doc, "PHẦN E — Đặt vào level và test", 1)
    add_step(doc, 8, "Đặt BP_CameraRig vào level", [
        "Content Browser → kéo BP_CameraRig vào Viewport.",
        "Đặt gần BP_StateManager hoặc trên conveyor.",
        "Ctrl+S lưu level.",
    ])

    add_step(doc, 9, "Test", [
        "Chạy python camera_server.py.",
        "Bấm Play.",
        "Bấm phím 3 → đèn phải xanh lá (Detected_OK).",
        "Bấm phím 6 → đèn phải đỏ cam (Fire_Alert), beam ẩn.",
        "Bấm phím 2 → đèn vàng (Scanning), beam hiện.",
    ])

    set_heading(doc, "Checklist hoàn thành TASK 2", 1)
    for item in [
        "BP_CameraRig có components: CameraBody, ScanLight, ScanBeam.",
        "Hàm GetColorForState trả đúng màu cho 6 state.",
        "Hàm ApplyState: đổi màu đèn + hiện/ẩn ScanBeam.",
        "BeginPlay bind OnStateChanged từ BP_StateManager.",
        "ScanBeam chỉ hiện khi state = Scanning.",
        "Đèn đổi màu ngay khi bấm phím 1–6.",
        "Compile không lỗi.",
    ]:
        doc.add_paragraph(f"☐  {item}", style="List Bullet")

    doc.save(os.path.join(OUTPUT_DIR, "TASK2_BP_CameraRig.docx"))
    print("✓ TASK2_BP_CameraRig.docx")


# ─────────────────────────────────────────────────────────
# TASK 3: WBP_ScanResult HUD
# ─────────────────────────────────────────────────────────
def make_task3():
    doc = Document()
    header_page(doc, 3, "WBP_ScanResult HUD",
                "Widget UMG | object_id / type / weight / temp / status | Counter OK/Warning/Error | Red overlay Fire_Alert")

    set_heading(doc, "Tổng quan", 1)
    doc.add_paragraph(
        "WBP_ScanResult là Widget Blueprint (UMG) hiển thị thông tin object đang được quét "
        "và bộ đếm trạng thái. Khi state là Fire_Alert, toàn bộ màn hình có overlay đỏ nhấp nháy."
    )

    set_heading(doc, "Yêu cầu trước khi làm", 2)
    for item in [
        "TASK 1 hoàn thành: BP_StateManager có OnStateChanged và các biến Count*.",
        "BP_StateManager đã có biến LastObjectId, LastStatus, LastTemperature.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    set_heading(doc, "PHẦN A — Tạo Widget Blueprint", 1)
    add_step(doc, 1, "Tạo WBP_ScanResult", [
        "Content Browser → Content/CameraModule/UI.",
        "Right-click → User Interface → Widget Blueprint.",
        "Đặt tên: WBP_ScanResult.",
        "Double-click mở UMG Designer.",
    ])

    set_heading(doc, "PHẦN B — Thiết kế layout", 1)
    add_step(doc, 2, "Thêm Canvas Panel gốc", [
        "Palette bên trái → kéo Canvas Panel vào graph.",
        "Size: Full screen (hoặc set trong Anchors: full preset).",
    ])

    add_step(doc, 3, "Thêm panel thông tin (góc trên trái)", [
        "Kéo Vertical Box vào Canvas, đặt anchor góc trên trái.",
        "Trong Vertical Box, thêm lần lượt các Text block:",
        "    Tên: txt_ObjectId    |  Text: 'Object: ---'",
        "    Tên: txt_Type        |  Text: 'Type: ---'",
        "    Tên: txt_Weight      |  Text: 'Weight: ---'",
        "    Tên: txt_Temperature |  Text: 'Temp: ---'",
        "    Tên: txt_Status      |  Text: 'Status: ---'",
        "    Tên: txt_State       |  Text: 'State: Idle'",
        "Font size 14–16, màu trắng, nền semi-transparent (dùng Border/Image làm background).",
    ])

    add_step(doc, 4, "Thêm panel counter (góc trên phải)", [
        "Kéo Vertical Box thứ hai vào Canvas, anchor góc trên phải.",
        "Thêm Text blocks:",
        "    Tên: txt_CountOK      |  Text: 'OK: 0'       |  màu xanh lá",
        "    Tên: txt_CountWarning |  Text: 'Warning: 0'  |  màu cam",
        "    Tên: txt_CountError   |  Text: 'Error: 0'    |  màu đỏ",
    ])

    add_step(doc, 5, "Thêm Fire Alert Overlay (toàn màn hình)", [
        "Kéo Image vào Canvas → đặt anchor full screen.",
        "Tên: img_FireOverlay.",
        "Tint color: đỏ bán trong suốt (R=1, G=0, B=0, A=0.4).",
        "Visibility: Hidden (mặc định ẩn).",
        "Đặt Z-Order cao nhất để nằm trên cùng.",
    ])

    set_heading(doc, "PHẦN C — Bind dữ liệu vào Widget", 1)
    add_step(doc, 6, "Tạo biến tham chiếu BP_StateManager", [
        "My Blueprint (trong Widget) → VARIABLES → +.",
        "Tên: StateManagerRef  |  Kiểu: BP_StateManager (Object Reference).",
        "Instance Editable: bật (để set từ ngoài vào).",
    ])

    add_step(doc, 7, "Tạo function RefreshUI", [
        "My Blueprint → FUNCTIONS → + → đặt tên: RefreshUI.",
        "Mở graph.",
        "─── Lấy dữ liệu từ StateManagerRef ───",
        "Get StateManagerRef → Get LastObjectId → Format Text: 'Object: {id}'.",
        "→ Set Text của txt_ObjectId.",
        "Làm tương tự cho LastStatus, LastTemperature (Format: 'Temp: {val}°C').",
        "Get CurrentState → Convert Enum to String → Set txt_State.",
        "Get CountOK → Convert to String → 'OK: {n}' → Set txt_CountOK.",
        "Get CountWarning → 'Warning: {n}' → Set txt_CountWarning.",
        "Get CountError → 'Error: {n}' → Set txt_CountError.",
        "─── Fire Alert Overlay ───",
        "Get CurrentState → Equal == Fire_Alert → Branch.",
        "True: Set Visibility của img_FireOverlay → Visible.",
        "False: Set Visibility → Hidden.",
        "Compile + Save.",
    ])

    add_step(doc, 8, "Bind OnStateChanged để tự động refresh", [
        "Mở Event Graph của Widget.",
        "Tạo Custom Event: OnConstruct (hoặc dùng Event Construct sẵn có).",
        "Trong Event Construct:",
        "    Get All Actors Of Class (BP_StateManager) → Get [0].",
        "    Set StateManagerRef = actor đó.",
        "    Bind Event to OnStateChanged (của StateManagerRef).",
        "    Tạo Custom Event: HandleStateChanged (input NewState EModuleState).",
        "    Trong HandleStateChanged: gọi RefreshUI.",
        "Compile + Save.",
    ])

    set_heading(doc, "PHẦN D — Thêm Widget vào màn hình trong level", 1)
    add_step(doc, 9, "Tạo HUD class hoặc thêm Widget trong Level Blueprint", [
        "Cách đơn giản: mở Level Blueprint (Blueprints → Open Level Blueprint).",
        "Event BeginPlay → Create Widget (Class = WBP_ScanResult) → Add to Viewport.",
        "Compile + Save Level Blueprint.",
    ])

    add_note(doc, "Có thể tạo BP_HUD class riêng kế thừa HUD và add widget trong đó "
                  "nếu muốn cấu trúc rõ ràng hơn.")

    set_heading(doc, "PHẦN E — Test", 1)
    add_step(doc, 10, "Chạy test toàn bộ", [
        "python camera_server.py đang chạy.",
        "Bấm Play.",
        "HUD hiện góc trên trái: thông tin object từ mock data.",
        "Counter OK/Warning/Error tăng theo poll.",
        "Bấm phím 6 (Fire_Alert): overlay đỏ bao phủ màn hình.",
        "Bấm phím 1 (Idle): overlay biến mất.",
    ])

    set_heading(doc, "Checklist hoàn thành TASK 3", 1)
    for item in [
        "WBP_ScanResult tồn tại trong Content Browser.",
        "Hiển thị: object_id, type, weight, temperature, status, state.",
        "Counter: OK / Warning / Error.",
        "img_FireOverlay hiện khi Fire_Alert, ẩn khi state khác.",
        "RefreshUI được gọi mỗi khi OnStateChanged fire.",
        "Widget xuất hiện trên màn hình khi Play.",
        "Compile không lỗi.",
    ]:
        doc.add_paragraph(f"☐  {item}", style="List Bullet")

    doc.save(os.path.join(OUTPUT_DIR, "TASK3_WBP_ScanResult_HUD.docx"))
    print("✓ TASK3_WBP_ScanResult_HUD.docx")


# ─────────────────────────────────────────────────────────
# TASK 4: Main Level Map
# ─────────────────────────────────────────────────────────
def make_task4():
    doc = Document()
    header_page(doc, 4, "Main Level Map",
                "Đặt conveyor + BP_CameraRig + BP_StateManager | Ánh sáng môi trường | Play → state tự cycling")

    set_heading(doc, "Tổng quan", 1)
    doc.add_paragraph(
        "TASK 4 là giai đoạn tích hợp: ghép tất cả Blueprint đã tạo vào một level hoàn chỉnh, "
        "thiết lập ánh sáng và môi trường, sau đó chạy Play để kiểm tra toàn bộ flow "
        "Camera Scanning từ mock data → state → đèn → HUD."
    )

    set_heading(doc, "Yêu cầu trước khi làm", 2)
    for item in [
        "TASK 1 hoàn thành: BP_StateManager.",
        "TASK 2 hoàn thành: BP_CameraRig.",
        "TASK 3 hoàn thành: WBP_ScanResult + Level Blueprint add widget.",
        "Python server camera_server.py đã test chạy được.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    set_heading(doc, "PHẦN A — Tạo hoặc cấu hình Main Level", 1)
    add_step(doc, 1, "Tạo map mới (nếu chưa có map chính)", [
        "File → New Level → chọn Basic (có sẵn sàn + ánh sáng) hoặc Empty Level.",
        "Save As → Content/CameraModule/Maps → đặt tên: L_CameraScan_Main.",
        "Ctrl+S.",
    ])

    add_step(doc, 2, "Cấu hình ánh sáng môi trường (nếu dùng Empty Level)", [
        "Place Actors → Light → Directional Light → kéo vào scene.",
        "Place Actors → Light → Sky Light → kéo vào scene.",
        "Place Actors → Visual Effects → Sky Atmosphere (tùy chọn, đẹp hơn).",
        "Ctrl+S.",
    ])

    set_heading(doc, "PHẦN B — Đặt các Actor vào level", 1)
    add_step(doc, 3, "Đặt Conveyor belt (sàn/belt tạm nếu chưa có mesh riêng)", [
        "Tạm thời dùng StaticMesh cube scale dài (conveyor giả):",
        "Place Actors → Shape → Cube → đặt vào giữa scene.",
        "Scale: (5.0, 1.0, 0.1) để tạo băng tải dài.",
        "Material: M_Metal hoặc màu xám công nghiệp.",
        "(Sau khi có mesh thật) xóa cube và thay bằng blueprint conveyor.",
    ])

    add_step(doc, 4, "Đặt BP_StateManager", [
        "Content Browser → Content/CameraModule/Core → kéo BP_StateManager vào Viewport.",
        "Đặt gần conveyor.",
        "Chọn actor → Details → Auto Receive Input → Player 0.",
        "Kiểm tra Http Base Url = http://127.0.0.1:8080 (nếu dùng bản C++).",
    ])

    add_step(doc, 5, "Đặt BP_CameraRig", [
        "Content Browser → Content/CameraModule/Actors → kéo BP_CameraRig vào Viewport.",
        "Đặt phía trên conveyor (như camera treo trần).",
        "Rotation: chỉnh để đèn SpotLight hướng xuống conveyor.",
        "Ctrl+D để duplicate nếu muốn 2 camera rig.",
    ])

    add_step(doc, 6, "Xác nhận Widget được add (từ Level Blueprint)", [
        "Blueprints (toolbar trên) → Open Level Blueprint.",
        "Kiểm tra Event BeginPlay đã có: Create Widget WBP_ScanResult → Add to Viewport.",
        "Nếu chưa có: thêm vào.",
        "Compile + Save Level Blueprint.",
    ])

    set_heading(doc, "PHẦN C — Cài đặt Camera và góc nhìn", 1)
    add_step(doc, 7, "Đặt Player Start và điều chỉnh góc nhìn", [
        "Place Actors → tìm Player Start → kéo vào scene.",
        "Đặt ở vị trí nhìn bao quát conveyor và BP_CameraRig.",
        "(Tùy chọn) Tạo BP_TopDownCamera hoặc dùng Spectator Pawn để xoay nhìn thoải mái.",
    ])

    set_heading(doc, "PHẦN D — Kiểm tra toàn bộ flow", 1)
    add_step(doc, 8, "Chạy server và Play", [
        "Mở terminal: python camera_server.py.",
        "Trong UE: bấm Play.",
        "Quan sát:",
        "    1. Đèn BP_CameraRig đổi màu theo state.",
        "    2. ScanBeam hiện khi state = Scanning.",
        "    3. HUD góc trái: thông tin object + counter tăng.",
        "    4. HUD hiện overlay đỏ khi bấm phím 6 (Fire_Alert).",
        "Bấm lần lượt phím 1–6 kiểm tra toàn bộ state.",
    ])

    add_step(doc, 9, "Quay video demo", [
        "Sử dụng OBS Studio, ShareX, hoặc Windows Game Bar (Win+G).",
        "Quay ít nhất 60 giây bao gồm:",
        "    - Auto polling từ server (state tự đổi).",
        "    - Thao tác phím 1–6.",
        "    - HUD hiển thị thông tin.",
        "    - Hiệu ứng Fire_Alert.",
        "Lưu video vào thư mục DEMO/ trong project.",
    ])

    set_heading(doc, "PHẦN E — Checklist tích hợp cuối", 1)
    for item in [
        "Level L_CameraScan_Main tồn tại và mở được.",
        "BP_StateManager trong level, Auto Receive Input = Player 0.",
        "BP_CameraRig trong level, đèn đổi màu theo state.",
        "WBP_ScanResult hiện trên màn hình khi Play.",
        "Phím 1–6 force state hoạt động.",
        "State tự đổi theo mock data mỗi ~3 giây.",
        "ScanBeam chỉ hiện khi Scanning.",
        "Fire_Alert hiện overlay đỏ.",
        "Video demo đã quay và lưu.",
        "Compile + Save toàn bộ không lỗi.",
    ]:
        doc.add_paragraph(f"☐  {item}", style="List Bullet")

    set_heading(doc, "Chuẩn bị nộp bài", 1)
    doc.add_paragraph("Sau khi hoàn thành TASK 4, chuẩn bị các mục sau để nộp:")
    for item in [
        "Báo cáo (.docx / .pdf): bối cảnh, mục tiêu, phạm vi, mô tả module, thiết kế, mock data, kết quả, hạn chế, hướng phát triển.",
        "Video demo (.mp4): tối thiểu 60 giây, có âm thanh hoặc caption.",
        "Source code: commit toàn bộ lên Git, README.md đầy đủ.",
        "Slide bảo vệ (.pptx): giới thiệu, kiến trúc, chức năng, demo, khó khăn, hướng phát triển.",
        "Hồ sơ kỹ thuật: tên module, input/output, luồng dữ liệu mock, kiến trúc, cách thể hiện state.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(os.path.join(OUTPUT_DIR, "TASK4_MainLevel_Map.docx"))
    print("✓ TASK4_MainLevel_Map.docx")


if __name__ == "__main__":
    make_task1()
    make_task2()
    make_task3()
    make_task4()
    print(f"\nTất cả file đã được tạo tại: {OUTPUT_DIR}")
